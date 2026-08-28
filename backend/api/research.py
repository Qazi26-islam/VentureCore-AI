import logging
import json
import io
import threading
import re
import csv
import html
from datetime import datetime
from typing import Optional
from fastapi import APIRouter, HTTPException, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse, Response
import markdown as markdown_lib
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment

from backend.models.schemas import (
    ResearchRequest, StartResponse, StatusResponse,
    HistoryItem, JobDetailResponse, MessageItem,
    FollowUpRequest, FollowUpResponse,
    RenameRequest, FavoriteRequest,
    CompanyProfileRequest, CompanyProfileResponse,
    DataQualityResponse,
    SupplierRequest, SupplierItem, ProductRequest, StockMovementRequest, InventoryItem,
    InventoryDashboardResponse,
    InventoryQuestionRequest, InventoryQuestionResponse,
    CustomerRequest, CustomerItem, SaleRequest, SaleRecord, SalesDashboardResponse,
    SalesQuestionRequest, SalesQuestionResponse,
    FinanceTransactionRequest, FinanceTransactionItem, FinanceDashboardResponse,
    FinanceQuestionRequest, FinanceQuestionResponse,
)
from backend.agents.coordinator import run_research
from backend.agents import followup, inventory as inventory_agent, sales as sales_agent
from backend.agents import finance_operations as finance_agent
from backend import jobs
from backend.db import (
    DEFAULT_ORGANIZATION_ID,
    DEMO_ORGANIZATION_ID,
    get_connection,
    get_demo_user_id,
)
from backend.tools import ToolContext, invoke_tool
from backend.observability import dashboard_data, instrument_client, run_traced_agent

router = APIRouter()
logger = logging.getLogger("research_api")


def require_admin(request: Request) -> int:
    user_id = request.session.get("user_id")
    if user_id is None:
        raise HTTPException(status_code=401, detail="Administrator sign-in required.")
    conn = get_connection()
    row = conn.execute("SELECT role FROM users WHERE id = ?", (user_id,)).fetchone()
    conn.close()
    if row is None or row["role"] != "admin":
        raise HTTPException(status_code=403, detail="Administrator access required.")
    return int(user_id)


@router.get("/internal/agent-runs", response_class=HTMLResponse)
def agent_runs_page(request: Request) -> HTMLResponse:
    require_admin(request)
    data = dashboard_data()
    run_rows = "".join(
        "<tr>"
        f"<td>{html.escape(str(row['created_at']))}</td>"
        f"<td>{html.escape(str(row['agent_name']))}</td>"
        f"<td>{html.escape(str(row['status']))}</td>"
        f"<td>{int(row['latency_ms'] or 0)} ms</td>"
        f"<td>{int(row['input_tokens'])} / {int(row['output_tokens'])}</td>"
        f"<td>${int(row['cost_minor']) / 100:.2f}</td>"
        f"<td>{html.escape(str(row['trigger_text'] or ''))}</td>"
        "</tr>"
        for row in data["runs"]
    ) or "<tr><td colspan='7'>No agent runs recorded yet.</td></tr>"
    tool_rows = "".join(
        f"<li>{html.escape(str(item['tool_name']))}: {int(item['calls'])}</li>"
        for item in data["tool_frequency"]
    ) or "<li>No tool calls recorded.</li>"
    document = f"""<!doctype html><html><head><meta charset='utf-8'>
    <title>VentureCore agent runs</title><style>
    body{{font:14px system-ui;background:#0d1117;color:#e6edf3;padding:24px}}
    table{{width:100%;border-collapse:collapse}}th,td{{padding:10px;border-bottom:1px solid #30363d;text-align:left}}
    .metrics{{display:flex;gap:24px;margin:20px 0}}.metric{{background:#161b22;padding:16px;border-radius:8px}}
    </style></head><body><h1>Agent runs</h1>
    <div class='metrics'><div class='metric'>Median latency<br><strong>{data['median_latency_ms']} ms</strong></div>
    <div class='metric'>p95 latency<br><strong>{data['p95_latency_ms']} ms</strong></div>
    <div class='metric'>Error rate<br><strong>{data['error_rate']:.1f}%</strong></div></div>
    <h2>Tool-call frequency</h2><ul>{tool_rows}</ul><h2>Recent runs</h2>
    <table><thead><tr><th>Started</th><th>Agent</th><th>Status</th><th>Latency</th>
    <th>Tokens in/out</th><th>Cost</th><th>Trigger</th></tr></thead><tbody>{run_rows}</tbody></table>
    </body></html>"""
    return HTMLResponse(document)


def get_user_id(request: Request):
    user_id = request.session.get("user_id")
    if user_id is not None:
        return user_id
    if is_demo_request(request):
        return get_demo_user_id()
    return None


def is_demo_request(request: Request) -> bool:
    return bool(request.session.get("demo_mode"))


def get_organization_id(request: Request) -> int:
    return DEMO_ORGANIZATION_ID if is_demo_request(request) else DEFAULT_ORGANIZATION_ID


def require_login(request: Request) -> int:
    user_id = get_user_id(request)
    if user_id is None:
        raise HTTPException(status_code=401, detail="You must be logged in.")
    return user_id


def require_write_access(request: Request) -> int:
    if is_demo_request(request):
        raise HTTPException(
            status_code=403,
            detail="The public demo is read-only. Sign up to make changes in your own workspace.",
        )
    return require_login(request)


def reject_demo_write(request: Request) -> None:
    if is_demo_request(request):
        raise HTTPException(
            status_code=403,
            detail="The public demo is read-only. Sign up to make changes in your own workspace.",
        )


@router.post("/demo/start")
def start_demo(request: Request):
    request.session.pop("user_id", None)
    request.session["demo_mode"] = True
    return {"status": "demo", "organization_id": DEMO_ORGANIZATION_ID}


@router.post("/demo/exit")
def exit_demo(request: Request):
    request.session.pop("demo_mode", None)
    return {"status": "exited"}


def _profile_to_context(profile) -> str:
    if profile is None:
        return ""
    fields = [
        ("Company", profile["company_name"]),
        ("Industry", profile["industry"]),
        ("Country / market", profile["country"]),
        ("Currency", profile["currency"]),
        ("Products / services", profile["products_services"]),
        ("Target customers", profile["target_customers"]),
        ("Main competitors", profile["main_competitors"]),
        ("Monthly budget (minor units)", profile["monthly_budget_minor"]),
        ("Business goals", profile["business_goals"]),
        ("Business stage", profile["business_stage"]),
    ]
    lines = [f"{label}: {value}" for label, value in fields if value]
    return "Saved company profile:\n" + "\n".join(lines)


@router.get("/company/profile", response_model=Optional[CompanyProfileResponse])
def get_company_profile(request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    row = conn.execute(
        "SELECT * FROM company_profiles WHERE user_id = ? AND organization_id = ?",
        (user_id, organization_id),
    ).fetchone()
    conn.close()
    return dict(row) if row else None


@router.put("/company/profile", response_model=CompanyProfileResponse)
def save_company_profile(body: CompanyProfileRequest, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    values = body.model_dump()
    conn = get_connection()
    conn.execute(
        """INSERT INTO company_profiles
           (user_id, organization_id, company_name, industry, country, currency, products_services,
            target_customers, main_competitors, monthly_budget_minor, business_goals, business_stage,
            source)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'manual')
           ON CONFLICT(organization_id, user_id) DO UPDATE SET
             organization_id=excluded.organization_id, company_name=excluded.company_name,
             industry=excluded.industry, country=excluded.country,
             currency=excluded.currency, products_services=excluded.products_services,
             target_customers=excluded.target_customers, main_competitors=excluded.main_competitors,
             monthly_budget_minor=excluded.monthly_budget_minor, business_goals=excluded.business_goals,
             business_stage=excluded.business_stage, updated_at=CURRENT_TIMESTAMP""",
        (user_id, organization_id, *values.values()),
    )
    conn.commit()
    row = conn.execute(
        "SELECT * FROM company_profiles WHERE user_id = ? AND organization_id = ?",
        (user_id, organization_id),
    ).fetchone()
    conn.close()
    return dict(row)


def _is_blank(value) -> bool:
    return value is None or (isinstance(value, str) and not value.strip())


def _looks_numeric_column(name: str) -> bool:
    return any(word in name.lower() for word in ("sales", "revenue", "cost", "price", "stock", "quantity", "qty", "expense", "amount", "margin"))


@router.post("/company/data/quality", response_model=DataQualityResponse)
async def check_data_quality(request: Request, file: UploadFile = File(...), data_type: str = Form("Business data")):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    filename = file.filename or "upload"
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if extension not in {"csv", "xlsx"}:
        raise HTTPException(status_code=400, detail="Upload a CSV or XLSX file.")
    raw = await file.read()
    if not raw or len(raw) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File must be between 1 byte and 10 MB.")
    try:
        if extension == "csv":
            decoded = raw.decode("utf-8-sig")
            rows = list(csv.reader(decoded.splitlines()))
        else:
            workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            sheet = workbook.active
            rows = [list(row) for row in sheet.iter_rows(values_only=True)]
            workbook.close()
    except Exception:
        raise HTTPException(status_code=400, detail="This file could not be read. Check that it is a valid CSV or XLSX file.")
    if len(rows) < 2:
        raise HTTPException(status_code=400, detail="Your file needs a header row and at least one data row.")
    columns = [str(cell).strip() if cell is not None else f"Column {index + 1}" for index, cell in enumerate(rows[0])]
    data_rows = rows[1:]
    missing_values = sum(1 for row in data_rows for cell in row if _is_blank(cell))
    normalized_rows = [tuple("" if _is_blank(cell) else str(cell).strip() for cell in row) for row in data_rows]
    duplicate_rows = len(normalized_rows) - len(set(normalized_rows))
    invalid_numeric_values = 0
    for row in data_rows:
        for index, cell in enumerate(row[:len(columns)]):
            if _looks_numeric_column(columns[index]) and not _is_blank(cell):
                try:
                    float(str(cell).replace(",", "").replace("RM", "").replace("BDT", "").replace("$", "").strip())
                except ValueError:
                    invalid_numeric_values += 1
    warnings = []
    if missing_values:
        warnings.append(f"{missing_values} missing value(s) found.")
    if duplicate_rows:
        warnings.append(f"{duplicate_rows} duplicate row(s) found.")
    if invalid_numeric_values:
        warnings.append(f"{invalid_numeric_values} value(s) in numeric-looking columns could not be read as numbers.")
    if not warnings:
        warnings.append("No common data-quality issues were found in this file.")
    result = DataQualityResponse(filename=filename, data_type=data_type, row_count=len(data_rows), column_count=len(columns), columns=columns, missing_values=missing_values, duplicate_rows=duplicate_rows, invalid_numeric_values=invalid_numeric_values, warnings=warnings)
    conn = get_connection()
    conn.execute(
        """INSERT INTO data_uploads
           (user_id, organization_id, filename, data_type, row_count, column_count, quality_summary, source)
           VALUES (?, ?, ?, ?, ?, ?, ?, 'csv_import')""",
        (user_id, organization_id, filename, data_type, result.row_count, result.column_count,
         json.dumps(result.model_dump())),
    )
    conn.commit()
    conn.close()
    return result


@router.post("/inventory/suppliers")
def create_supplier(body: SupplierRequest, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    cursor = conn.execute(
        """INSERT INTO suppliers
           (user_id, organization_id, name, contact_name, email, phone, lead_time_days, payment_terms, source)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'manual')""",
        (user_id, organization_id, body.name, body.contact_name, body.email, body.phone,
         body.lead_time_days, body.payment_terms),
    )
    conn.commit()
    supplier_id = cursor.lastrowid
    conn.close()
    return {"id": supplier_id, "status": "created"}


@router.get("/inventory/suppliers", response_model=list[SupplierItem])
def list_suppliers(request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    rows = conn.execute(
        """SELECT id, name, contact_name, email, phone, lead_time_days, payment_terms
           FROM suppliers WHERE user_id = ? AND organization_id = ? ORDER BY name COLLATE NOCASE""",
        (user_id, organization_id),
    ).fetchall()
    conn.close()
    return [SupplierItem(**dict(row)) for row in rows]


@router.post("/inventory/products")
def create_product(body: ProductRequest, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    if body.supplier_id is not None:
        supplier = conn.execute(
            "SELECT id FROM suppliers WHERE id = ? AND user_id = ? AND organization_id = ?",
            (body.supplier_id, user_id, organization_id),
        ).fetchone()
        if supplier is None:
            conn.close()
            raise HTTPException(status_code=404, detail="Supplier not found")
    try:
        cursor = conn.execute(
            """INSERT INTO products
               (user_id, organization_id, supplier_id, sku, name, category, unit_cost_minor,
                selling_price_minor, currency, reorder_point, lead_time_days, source)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'manual')""",
            (user_id, organization_id, body.supplier_id, body.sku, body.name, body.category,
             body.unit_cost_minor, body.selling_price_minor, body.currency, body.reorder_point,
             body.lead_time_days),
        )
        conn.commit()
        product_id = cursor.lastrowid
    except Exception as exc:
        conn.close()
        if "UNIQUE" in str(exc):
            raise HTTPException(status_code=400, detail="This SKU already exists in your company workspace.")
        raise
    conn.close()
    return {"id": product_id, "status": "created"}


@router.post("/inventory/products/{product_id}/movement")
def record_stock_movement(product_id: int, body: StockMovementRequest, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    product = conn.execute(
        "SELECT id, currency FROM products WHERE id = ? AND user_id = ? AND organization_id = ?",
        (product_id, user_id, organization_id),
    ).fetchone()
    if product is None:
        conn.close()
        raise HTTPException(status_code=404, detail="Product not found")
    direction = 1 if body.transaction_type in {"received", "adjustment"} else -1
    quantity_change = body.quantity * direction
    current_stock = conn.execute(
        """SELECT COALESCE(SUM(quantity_change), 0) FROM inventory_transactions
           WHERE product_id = ? AND organization_id = ?""",
        (product_id, organization_id),
    ).fetchone()[0]
    if current_stock + quantity_change < 0:
        conn.close()
        raise HTTPException(status_code=400, detail=f"Not enough stock. Current stock is {current_stock}.")
    cursor = conn.execute(
        """INSERT INTO inventory_transactions
           (product_id, user_id, organization_id, transaction_type, quantity_change,
            unit_cost_minor, currency, reference_note, source)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'manual')""",
        (product_id, user_id, organization_id, body.transaction_type, quantity_change,
         body.unit_cost_minor, product["currency"], body.reference_note),
    )
    conn.commit()
    conn.close()
    return {"id": cursor.lastrowid, "current_stock": current_stock + quantity_change, "status": "recorded"}


@router.get("/inventory/dashboard", response_model=InventoryDashboardResponse)
def inventory_dashboard(request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    result = invoke_tool(
        "get_inventory_snapshot",
        ToolContext(organization_id=organization_id, user_id=user_id),
        {"lookback_days": 30},
    )
    if not result.ok:
        raise HTTPException(status_code=500, detail="Could not calculate inventory metrics.")
    data = result.data or {}
    dashboard = data["dashboard"]
    items = [
        InventoryItem(
            **{
                **item,
                "units_sold_30d": item["units_sold"],
            }
        )
        for item in data["items"]
    ]
    return InventoryDashboardResponse(**dashboard, items=items)


@router.post("/inventory/ask", response_model=InventoryQuestionResponse)
def ask_inventory_agent(body: InventoryQuestionRequest, request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    try:
        answer = inventory_agent.run(
            question=body.question,
            organization_id=organization_id,
            user_id=user_id,
        )
    except Exception as exc:
        logger.exception("Inventory Agent failed: %s", exc)
        raise HTTPException(status_code=502, detail="The Inventory Agent is temporarily unavailable. Please try again.")
    return InventoryQuestionResponse(answer=answer)


@router.post("/sales/customers")
def create_customer(body: CustomerRequest, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    cursor = conn.execute(
        """INSERT INTO customers (user_id, organization_id, name, email, phone, segment, notes, source)
           VALUES (?, ?, ?, ?, ?, ?, ?, 'manual')""",
        (user_id, organization_id, body.name, body.email, body.phone, body.segment, body.notes),
    )
    conn.commit()
    customer_id = cursor.lastrowid
    conn.close()
    return {"id": customer_id, "status": "created"}


@router.get("/sales/customers", response_model=list[CustomerItem])
def list_customers(request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    rows = conn.execute(
        """SELECT id, name, email, phone, segment, notes
           FROM customers WHERE user_id = ? AND organization_id = ? ORDER BY name COLLATE NOCASE""",
        (user_id, organization_id),
    ).fetchall()
    conn.close()
    return [CustomerItem(**dict(row)) for row in rows]


@router.post("/sales/orders", response_model=SaleRecord)
def record_sale(body: SaleRequest, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    result = invoke_tool(
        "record_sale",
        ToolContext(organization_id=organization_id, user_id=user_id, source="manual"),
        body.model_dump(),
    )
    if not result.ok:
        error = result.error
        status_code = 404 if error and error.code.endswith("_not_found") else 400
        if error and error.code == "tool_execution_failed":
            status_code = 500
        raise HTTPException(status_code=status_code, detail=error.message if error else "Could not record sale.")
    sale = result.data or {}
    effective_status = (
        "overdue"
        if sale["payment_status"] == "due"
        and sale.get("due_date")
        and sale["due_date"] < datetime.now().strftime("%Y-%m-%d")
        else sale["payment_status"]
    )
    return SaleRecord(
        id=sale["id"], customer_name=sale.get("customer_name"), product_name=sale["product_name"],
        quantity=sale["quantity"], unit_price_minor=sale["unit_price_minor"],
        total_amount_minor=sale["total_amount_minor"], currency=sale["currency"],
        payment_status=effective_status, due_date=sale.get("due_date"), created_at=sale["created_at"],
    )


@router.post("/sales/orders/{sale_id}/mark-paid")
def mark_sale_paid(sale_id: int, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    sale = conn.execute(
        """SELECT id, total_amount_minor, currency, payment_status FROM sales_orders
           WHERE id = ? AND user_id = ? AND organization_id = ?""",
        (sale_id, user_id, organization_id),
    ).fetchone()
    if sale is None:
        conn.close()
        raise HTTPException(status_code=404, detail="Sale not found.")
    if sale["payment_status"] != "paid":
        conn.execute(
            "UPDATE sales_orders SET payment_status = 'paid' WHERE id = ? AND organization_id = ?",
            (sale_id, organization_id),
        )
        conn.execute(
            """INSERT OR IGNORE INTO finance_transactions
               (user_id, organization_id, transaction_type, amount_minor, currency, category,
                description, source, external_id, related_sale_id, transaction_date)
               VALUES (?, ?, 'income', ?, ?, 'Sales Revenue', ?, 'sale', ?, ?, date('now'))""",
            (user_id, organization_id, sale["total_amount_minor"], sale["currency"],
             f"Payment received for sale #{sale_id}", str(sale_id), sale_id),
        )
        conn.commit()
    conn.close()
    return {"status": "paid", "sale_id": sale_id}


@router.get("/sales/dashboard", response_model=SalesDashboardResponse)
def sales_dashboard(request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    result = invoke_tool(
        "get_sales_snapshot",
        ToolContext(organization_id=organization_id, user_id=user_id),
        {"lookback_days": 30},
    )
    if not result.ok:
        raise HTTPException(status_code=500, detail="Could not calculate sales metrics.")
    data = result.data or {}
    dashboard = data["dashboard"]
    return SalesDashboardResponse(
        revenue_30d_minor=dashboard["revenue_minor"],
        cash_collected_30d_minor=dashboard["cash_collected_minor"],
        outstanding_amount_minor=dashboard["outstanding_amount_minor"],
        overdue_receivables_minor=dashboard["overdue_receivables_minor"],
        currency=dashboard["currency"],
        orders_30d=dashboard["orders"],
        top_customer=dashboard["top_customer"],
        workings=dashboard["workings"],
        recent_sales=[SaleRecord(**row) for row in data["recent_sales"]],
    )


@router.post("/sales/ask", response_model=SalesQuestionResponse)
def ask_sales_agent(body: SalesQuestionRequest, request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    try:
        answer = sales_agent.run(body.question, organization_id=organization_id, user_id=user_id)
    except Exception as exc:
        logger.exception("Sales Agent failed: %s", exc)
        raise HTTPException(status_code=502, detail="The Sales Agent is temporarily unavailable. Please try again.")
    return SalesQuestionResponse(answer=answer)


@router.post("/finance/transactions", response_model=FinanceTransactionItem)
def create_finance_transaction(body: FinanceTransactionRequest, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    transaction_date = body.transaction_date or datetime.now().strftime("%Y-%m-%d")
    try:
        datetime.strptime(transaction_date, "%Y-%m-%d")
    except ValueError:
        raise HTTPException(status_code=400, detail="Transaction date must use YYYY-MM-DD format.")
    conn = get_connection()
    cursor = conn.execute(
        """INSERT INTO finance_transactions
           (user_id, organization_id, transaction_type, amount_minor, currency, category,
            description, source, transaction_date)
           VALUES (?, ?, ?, ?, ?, ?, ?, 'manual', ?)""",
        (user_id, organization_id, body.transaction_type, body.amount_minor, body.currency,
         body.category, body.description, transaction_date),
    )
    conn.commit()
    row = conn.execute(
        """SELECT id, transaction_type, amount_minor, currency, category, description, source,
                  transaction_date FROM finance_transactions
           WHERE id = ? AND organization_id = ?""",
        (cursor.lastrowid, organization_id),
    ).fetchone()
    conn.close()
    return FinanceTransactionItem(**dict(row))


@router.get("/finance/dashboard", response_model=FinanceDashboardResponse)
def finance_dashboard(request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    result = invoke_tool(
        "get_finance_snapshot",
        ToolContext(organization_id=organization_id, user_id=user_id),
        {"lookback_days": 30},
    )
    if not result.ok:
        raise HTTPException(status_code=500, detail="Could not calculate finance metrics.")
    data = result.data or {}
    dashboard = data["dashboard"]
    expense_workings = {
        item["category"]: {
            "tool": "get_finance_snapshot",
            "inputs": {"lookback_days": 30},
            "source_row_ids": item["source_row_ids"],
        }
        for item in data["expense_categories"]
    }
    return FinanceDashboardResponse(
        income_30d_minor=dashboard["income_minor"],
        expenses_30d_minor=dashboard["expenses_minor"],
        net_cash_flow_30d_minor=dashboard["net_cash_flow_minor"],
        cash_balance_minor=dashboard["recorded_cash_balance_minor"],
        receivables_minor=dashboard["receivables_minor"],
        currency=dashboard["currency"],
        expense_breakdown_30d_minor={
            item["category"]: item["amount_minor"] for item in data["expense_categories"]
        },
        workings=dashboard["workings"],
        expense_workings=expense_workings,
        recent_transactions=[FinanceTransactionItem(**row) for row in data["recent_transactions"]],
    )


@router.post("/finance/ask", response_model=FinanceQuestionResponse)
def ask_finance_agent(body: FinanceQuestionRequest, request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    try:
        answer = finance_agent.run(body.question, organization_id=organization_id, user_id=user_id)
    except Exception as exc:
        logger.exception("Finance Agent failed: %s", exc)
        raise HTTPException(status_code=502, detail="The Finance Agent is temporarily unavailable. Please try again.")
    return FinanceQuestionResponse(answer=answer)


@router.post("/research/start", response_model=StartResponse)
def start_research(request: Request, body: ResearchRequest) -> StartResponse:
    reject_demo_write(request)
    user_id = get_user_id(request)
    organization_id = get_organization_id(request)
    research_question = body.question
    if body.use_company_profile:
        if user_id is None:
            raise HTTPException(status_code=401, detail="Log in to use a saved company profile.")
        conn = get_connection()
        profile = conn.execute(
            "SELECT * FROM company_profiles WHERE user_id = ? AND organization_id = ?",
            (user_id, organization_id),
        ).fetchone()
        conn.close()
        if profile is None:
            raise HTTPException(status_code=400, detail="Create your Company Profile before using it in research.")
        research_question = _profile_to_context(profile) + "\n\nResearch request:\n" + body.question
    job_id = jobs.create_job(research_question, organization_id)

    if user_id is not None:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            """INSERT INTO research_jobs
               (id, user_id, organization_id, question, report, sections, source)
               VALUES (?, ?, ?, ?, ?, ?, 'manual')""",
            (job_id, user_id, organization_id, research_question, None, "{}"),
        )
        conn.commit()
        conn.close()

    def _background_run():
        try:
            run_research(
                research_question, job_id, mode=body.mode, depth=body.depth,
                organization_id=organization_id,
            )
            if user_id is not None:
                job = jobs.get_job(job_id, organization_id)
                conn2 = get_connection()
                cursor2 = conn2.cursor()
                cursor2.execute(
                    """UPDATE research_jobs SET report = ?, sections = ?
                       WHERE id = ? AND organization_id = ?""",
                    (job["report"], json.dumps(job["sections"]), job_id, organization_id),
                )
                conn2.commit()
                conn2.close()
        except Exception as e:
            logger.exception("Research pipeline crashed")
            jobs.fail_job(job_id, str(e))

    thread = threading.Thread(target=_background_run, daemon=True)
    thread.start()

    return StartResponse(job_id=job_id)


@router.get("/research/status/{job_id}", response_model=StatusResponse)
def get_status(job_id: str, request: Request) -> StatusResponse:
    job = jobs.get_job(job_id, get_organization_id(request))
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")

    return StatusResponse(
        status=job["status"],
        stage=job["stage"],
        sections=job["sections"],
        report=job["report"],
        error=job["error"],
    )


@router.get("/research/history", response_model=list[HistoryItem])
def get_history(request: Request, q: str = "", favorites_only: bool = False):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    cursor = conn.cursor()

    query = "SELECT id, question, title, favorite, created_at FROM research_jobs WHERE user_id = ? AND organization_id = ?"
    params = [user_id, organization_id]

    if q:
        query += " AND (question LIKE ? OR title LIKE ?)"
        like = f"%{q}%"
        params.extend([like, like])

    if favorites_only:
        query += " AND favorite = 1"

    query += " ORDER BY favorite DESC, created_at DESC"

    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()

    return [
        HistoryItem(
            id=r["id"],
            question=r["question"],
            title=r["title"],
            favorite=bool(r["favorite"]),
            created_at=r["created_at"],
        )
        for r in rows
    ]


@router.get("/research/job/{job_id}", response_model=JobDetailResponse)
def get_job_detail(job_id: str, request: Request):
    user_id = require_login(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT * FROM research_jobs WHERE id = ? AND user_id = ? AND organization_id = ?",
        (job_id, user_id, organization_id),
    )
    row = cursor.fetchone()
    if row is None:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")

    cursor.execute(
        """SELECT role, content FROM follow_up_messages
           WHERE job_id = ? AND organization_id = ? ORDER BY created_at ASC""",
        (job_id, organization_id),
    )
    messages = [MessageItem(role=m["role"], content=m["content"]) for m in cursor.fetchall()]
    conn.close()

    sections = json.loads(row["sections"]) if row["sections"] else {}

    return JobDetailResponse(
        id=row["id"],
        question=row["question"],
        title=row["title"],
        favorite=bool(row["favorite"]),
        report=row["report"],
        sections=sections,
        messages=messages,
    )


@router.put("/research/job/{job_id}/rename")
def rename_job(job_id: str, body: RenameRequest, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id FROM research_jobs WHERE id = ? AND user_id = ? AND organization_id = ?",
        (job_id, user_id, organization_id),
    )
    if cursor.fetchone() is None:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    cursor.execute(
        "UPDATE research_jobs SET title = ? WHERE id = ? AND organization_id = ?",
        (body.title, job_id, organization_id),
    )
    conn.commit()
    conn.close()
    return {"status": "renamed"}


@router.put("/research/job/{job_id}/favorite")
def favorite_job(job_id: str, body: FavoriteRequest, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id FROM research_jobs WHERE id = ? AND user_id = ? AND organization_id = ?",
        (job_id, user_id, organization_id),
    )
    if cursor.fetchone() is None:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    cursor.execute(
        "UPDATE research_jobs SET favorite = ? WHERE id = ? AND organization_id = ?",
        (1 if body.favorite else 0, job_id, organization_id),
    )
    conn.commit()
    conn.close()
    return {"status": "updated"}


@router.delete("/research/job/{job_id}")
def delete_job(job_id: str, request: Request):
    user_id = require_write_access(request)
    organization_id = get_organization_id(request)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id FROM research_jobs WHERE id = ? AND user_id = ? AND organization_id = ?",
        (job_id, user_id, organization_id),
    )
    if cursor.fetchone() is None:
        conn.close()
        raise HTTPException(status_code=404, detail="Job not found")
    cursor.execute(
        "DELETE FROM follow_up_messages WHERE job_id = ? AND organization_id = ?",
        (job_id, organization_id),
    )
    cursor.execute(
        "DELETE FROM research_jobs WHERE id = ? AND organization_id = ?",
        (job_id, organization_id),
    )
    conn.commit()
    conn.close()
    return {"status": "deleted"}


from backend.agents import opportunity_finder
from backend.models.schemas import OpportunityRequest, OpportunityResponse, OpportunityItem


@router.post("/research/opportunities", response_model=OpportunityResponse)
def find_opportunities(body: OpportunityRequest, request: Request):
    organization_id = get_organization_id(request)
    user_id = get_user_id(request) or 1
    opportunity_finder.client = instrument_client(opportunity_finder.client)
    items = run_traced_agent(
        "Opportunity Finder Agent", organization_id, body.query,
        lambda: opportunity_finder.run(
            body.query, organization_id=organization_id, user_id=user_id
        ),
    )
    return OpportunityResponse(items=[OpportunityItem(**item) for item in items])


def _get_report_and_question(job_id: str, user_id, organization_id: int):
    if user_id is not None:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            "SELECT * FROM research_jobs WHERE id = ? AND user_id = ? AND organization_id = ?",
            (job_id, user_id, organization_id),
        )
        row = cursor.fetchone()
        conn.close()
        if row is None:
            return None, None
        return row["question"], row["report"]
    else:
        job = jobs.get_job(job_id, organization_id)
        if job is None:
            return None, None
        return job["question"], job["report"]


def _get_research_data(job_id: str, user_id, organization_id: int):
    if user_id is not None:
        conn = get_connection()
        row = conn.execute(
            """SELECT question, report, sections FROM research_jobs
               WHERE id = ? AND user_id = ? AND organization_id = ?""",
            (job_id, user_id, organization_id),
        ).fetchone()
        conn.close()
        if row is None:
            return None, None, {}
        return row["question"], row["report"], json.loads(row["sections"] or "{}")
    job = jobs.get_job(job_id, organization_id)
    if job is None:
        return None, None, {}
    return job["question"], job["report"], job.get("sections", {})


def _clean_export_text(text: str) -> str:
    text = re.sub(
        r"^\s*(MARKET|COMPETITOR|FINANCIAL)_CHART_DATA:[^\n]*(?:\n(?!\s*(?:\*\*Sources|Sources:|---|\*\*[A-Za-z]))[^\n]*)*",
        "",
        text or "",
        flags=re.MULTILINE | re.IGNORECASE,
    )
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


def _safe_report_name(question: str) -> str:
    return "".join(c if c.isalnum() or c in " -_" else "" for c in question)[:50].strip() or "report"


@router.get("/research/job/{job_id}/pdf")
def export_pdf(job_id: str, request: Request):
    user_id = get_user_id(request)
    question, report = _get_report_and_question(job_id, user_id, get_organization_id(request))

    if question is None:
        raise HTTPException(status_code=404, detail="Job not found")
    if not report:
        raise HTTPException(status_code=400, detail="This report isn't ready yet.")

    body_html = markdown_lib.markdown(report, extensions=["tables"])
    prepared_date = datetime.now().strftime("%d %B %Y")

    html_content = f"""
    <html>
    <head>
    <style>
        body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11px; color: #1a1a1a; }}
        h1 {{ color: #5b7cfa; font-size: 20px; margin-bottom: 2px; }}
        h2 {{ color: #5b7cfa; font-size: 15px; margin-top: 18px; margin-bottom: 6px; }}
        h3 {{ color: #5b7cfa; font-size: 13px; margin-top: 14px; }}
        p {{ line-height: 1.5; margin: 6px 0; }}
        table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
        th, td {{ border: 1px solid #cccccc; padding: 6px 8px; text-align: left; font-size: 10px; }}
        th {{ background: #f0f0f0; font-weight: bold; }}
        hr {{ border: none; border-top: 1px solid #cccccc; margin: 16px 0; }}
        .meta {{ color: #666666; font-size: 10px; margin-bottom: 20px; }}
        a {{ color: #5b7cfa; }}
    </style>
    </head>
    <body>
        <h1>VentureCore AI</h1>
        <div class="meta">
            Business Intelligence Report<br/>
            {question}<br/>
            Prepared: {prepared_date}
        </div>
        {body_html}
    </body>
    </html>
    """

    buffer = io.BytesIO()
    pisa.CreatePDF(html_content, dest=buffer)
    pdf_bytes = buffer.getvalue()
    buffer.close()

    safe_name = "".join(c if c.isalnum() or c in " -_" else "" for c in question)[:50].strip() or "report"

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'},
    )


@router.get("/research/job/{job_id}/docx")
def export_docx(job_id: str, request: Request):
    question, report, _ = _get_research_data(
        job_id, get_user_id(request), get_organization_id(request)
    )
    if question is None:
        raise HTTPException(status_code=404, detail="Job not found")
    if not report:
        raise HTTPException(status_code=400, detail="This report isn't ready yet.")

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    title = document.add_heading("VENTURECORE AI", 0)
    title.runs[0].font.color.rgb = RGBColor(91, 124, 250)
    document.add_paragraph("Business Intelligence Report", style="Subtitle")
    document.add_paragraph(question)
    document.add_paragraph(f"Prepared: {datetime.now().strftime('%d %B %Y')}")
    document.add_page_break()

    for raw_line in _clean_export_text(report).splitlines():
        line = raw_line.strip()
        if not line or line == "---":
            continue
        if line.startswith("**") and line.endswith("**"):
            document.add_heading(line.strip("*").rstrip(":"), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            paragraph = document.add_paragraph(line.strip("|").replace("|", "  |  "))
            paragraph.style = document.styles["No Spacing"]
        else:
            document.add_paragraph(line.replace("**", ""))

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    buffer = io.BytesIO()
    document.save(buffer)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{_safe_report_name(question)}.docx"'},
    )


@router.get("/research/job/{job_id}/xlsx")
def export_xlsx(job_id: str, request: Request):
    question, report, sections = _get_research_data(
        job_id, get_user_id(request), get_organization_id(request)
    )
    if question is None:
        raise HTTPException(status_code=404, detail="Job not found")
    if not report:
        raise HTTPException(status_code=400, detail="This report isn't ready yet.")

    workbook = Workbook()
    workbook.remove(workbook.active)
    sheet_data = [
        ("Executive Report", report),
        ("Market Research", sections.get("market_research", "")),
        ("Competitors", sections.get("competitor_analysis", "")),
        ("Financial Analysis", sections.get("financial_analysis", "")),
    ]
    header_fill = PatternFill("solid", fgColor="5B7CFA")
    for sheet_name, content in sheet_data:
        sheet = workbook.create_sheet(sheet_name)
        sheet.append(["VENTURECORE AI", sheet_name])
        sheet.append(["Business", question])
        sheet.append(["Prepared", datetime.now().strftime("%d %B %Y")])
        sheet.append([])
        sheet.append(["Report content"])
        for line in _clean_export_text(content).splitlines():
            if line.strip():
                sheet.append([line.replace("**", "")])
        for cell in sheet[1]:
            cell.fill = header_fill
            cell.font = Font(color="FFFFFF", bold=True)
        sheet.column_dimensions["A"].width = 110
        sheet.column_dimensions["B"].width = 70
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    sources_sheet = workbook.create_sheet("Sources")
    sources_sheet.append(["Source", "URL", "Quality"])
    source_number = 1
    combined = "\n".join([report] + list(sections.values()))
    seen_urls = set()
    for label, url in re.findall(r"\[([^]]+)\]\((https?://[^)]+)\)", combined):
        if url in seen_urls:
            continue
        seen_urls.add(url)
        domain = url.lower()
        quality = "Official / Government" if any(token in domain for token in [".gov", ".edu", "worldbank.org", "who.int", "oecd.org"]) else "Industry / Company source"
        sources_sheet.append([f"[{source_number}] {label}", url, quality])
        source_number += 1
    for cell in sources_sheet[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
    sources_sheet.column_dimensions["A"].width = 50
    sources_sheet.column_dimensions["B"].width = 75
    sources_sheet.column_dimensions["C"].width = 25
    sources_sheet.freeze_panes = "A2"

    buffer = io.BytesIO()
    workbook.save(buffer)
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{_safe_report_name(question)}.xlsx"'},
    )


@router.post("/research/job/{job_id}/message", response_model=FollowUpResponse)
def send_follow_up(job_id: str, body: FollowUpRequest, request: Request):
    reject_demo_write(request)
    user_id = get_user_id(request)
    organization_id = get_organization_id(request)

    if user_id is not None:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            "SELECT * FROM research_jobs WHERE id = ? AND user_id = ? AND organization_id = ?",
            (job_id, user_id, organization_id),
        )
        row = cursor.fetchone()
        if row is None:
            conn.close()
            raise HTTPException(status_code=404, detail="Job not found")

        cursor.execute(
            """SELECT role, content FROM follow_up_messages
               WHERE job_id = ? AND organization_id = ? ORDER BY created_at ASC""",
            (job_id, organization_id),
        )
        history = [{"role": m["role"], "content": m["content"]} for m in cursor.fetchall()]
        question = row["question"]
        report = row["report"] or ""
    else:
        job = jobs.get_job(job_id, organization_id)
        if job is None:
            raise HTTPException(status_code=404, detail="Job not found")
        history = job["messages"]
        question = job["question"]
        report = job["report"] or ""

    try:
        followup.client = instrument_client(followup.client)
        reply = run_traced_agent(
            "Follow-up Agent", organization_id, body.message,
            lambda: followup.run(
                question=question, report=report, history=history,
                new_message=body.message,
            ),
            job_id,
        )
    except Exception as e:
        if user_id is not None:
            conn.close()
        logger.exception("Follow-up failed")
        raise HTTPException(status_code=502, detail="Follow-up failed. Please try again.") from e

    if user_id is not None:
        cursor.execute(
            """INSERT INTO follow_up_messages
               (job_id, organization_id, role, content, source) VALUES (?, ?, ?, ?, 'manual')""",
            (job_id, organization_id, "user", body.message),
        )
        cursor.execute(
            """INSERT INTO follow_up_messages
               (job_id, organization_id, role, content, source) VALUES (?, ?, ?, ?, 'manual')""",
            (job_id, organization_id, "assistant", reply),
        )
        conn.commit()
        conn.close()
    else:
        jobs.add_message(job_id, "user", body.message)
        jobs.add_message(job_id, "assistant", reply)

    return FollowUpResponse(reply=reply)
